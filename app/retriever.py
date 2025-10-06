"""
Document retriever module for GIS and Remote Sensing RAG system.
Handles document loading, text splitting, embedding, and retrieval.
"""

import os
import pickle
import logging
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
from datetime import datetime

import PyPDF2
import docx
import faiss
import numpy as np
import pandas as pd
from sentence_transformers import SentenceTransformer
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.vectorstores import FAISS
from unstructured.partition.auto import partition

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class GISDocumentRetriever:
    """
    A specialized document retriever for GIS and Remote Sensing documents.
    """

    def __init__(
        self,
        data_dir: str = "data",
        embeddings_dir: str = "embeddings",
        embedding_model: str = "sentence-transformers/all-MiniLM-L6-v2",
        chunk_size: int = 500,
        chunk_overlap: int = 100,
        use_faiss: bool = True
    ):
        """
        Initialize the GIS document retriever.

        Args:
            data_dir: Directory containing GIS documents
            embeddings_dir: Directory to store vector embeddings
            embedding_model: Name of the embedding model
            chunk_size: Size of text chunks for splitting
            chunk_overlap: Overlap between text chunks
            use_faiss: Whether to use FAISS for vector storage
        """
        self.data_dir = Path(data_dir)
        self.embeddings_dir = Path(embeddings_dir)
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.use_faiss = use_faiss

        # Create directories if they don't exist
        self.data_dir.mkdir(exist_ok=True)
        self.embeddings_dir.mkdir(exist_ok=True)

        # Initialize text splitter
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )

        # Initialize embedding model
        logger.info(f"Loading embedding model: {embedding_model}")
        self.embeddings = HuggingFaceEmbeddings(
            model_name=embedding_model,
            model_kwargs={'device': 'cpu'}
        )

        # Initialize vector store
        self.vector_store = None
        self.documents = []

    def load_pdf(self, file_path: Path) -> List[str]:
        """
        Extract text from PDF file.

        Args:
            file_path: Path to PDF file

        Returns:
            List of text chunks from the PDF
        """
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text_chunks = []

                for page_num, page in enumerate(pdf_reader.pages):
                    text = page.extract_text()
                    if text.strip():
                        text_chunks.append({
                            'text': text,
                            'source': file_path.name,
                            'page': page_num + 1,
                            'type': 'pdf'
                        })

                logger.info(f"Extracted {len(text_chunks)} pages from {file_path.name}")
                return text_chunks

        except Exception as e:
            logger.error(f"Error loading PDF {file_path}: {str(e)}")
            return []

    def load_docx(self, file_path: Path) -> List[str]:
        """
        Extract text from DOCX file.

        Args:
            file_path: Path to DOCX file

        Returns:
            List of text chunks from the DOCX
        """
        try:
            doc = docx.Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])

            if text.strip():
                logger.info(f"Extracted text from {file_path.name}")
                return [{
                    'text': text,
                    'source': file_path.name,
                    'page': 1,
                    'type': 'docx'
                }]
            return []

        except Exception as e:
            logger.error(f"Error loading DOCX {file_path}: {str(e)}")
            return []

    def load_txt(self, file_path: Path) -> List[str]:
        """
        Extract text from TXT file.

        Args:
            file_path: Path to TXT file

        Returns:
            List of text chunks from the TXT
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()

            if text.strip():
                logger.info(f"Extracted text from {file_path.name}")
                return [{
                    'text': text,
                    'source': file_path.name,
                    'page': 1,
                    'type': 'txt'
                }]
            return []

        except Exception as e:
            logger.error(f"Error loading TXT {file_path}: {str(e)}")
            return []

    def load_with_unstructured(self, file_path: Path) -> List[Dict[str, Any]]:
        """
        Extract text using unstructured library for various file types.

        Args:
            file_path: Path to file

        Returns:
            List of text chunks with metadata
        """
        try:
            elements = partition(filename=str(file_path))
            text_chunks = []

            for element in elements:
                if hasattr(element, 'text') and element.text.strip():
                    text_chunks.append({
                        'text': element.text,
                        'source': file_path.name,
                        'page': getattr(element, 'page_number', 1),
                        'type': file_path.suffix.lower().lstrip('.'),
                        'element_type': str(type(element).__name__)
                    })

            logger.info(f"Extracted {len(text_chunks)} elements from {file_path.name}")
            return text_chunks

        except Exception as e:
            logger.error(f"Error loading {file_path} with unstructured: {str(e)}")
            return []

    def load_documents(self) -> List[Document]:
        """
        Load all documents from the data directory.

        Returns:
            List of Document objects
        """
        logger.info(f"Loading documents from {self.data_dir}")

        all_text_chunks = []

        # Supported file extensions
        supported_extensions = {'.pdf', '.docx', '.txt', '.md', '.rtf'}

        for file_path in self.data_dir.rglob('*'):
            if file_path.is_file() and file_path.suffix.lower() in supported_extensions:
                logger.info(f"Processing file: {file_path}")

                # Try different loading methods based on file type
                if file_path.suffix.lower() == '.pdf':
                    chunks = self.load_pdf(file_path)
                elif file_path.suffix.lower() == '.docx':
                    chunks = self.load_docx(file_path)
                elif file_path.suffix.lower() in ['.txt', '.md']:
                    chunks = self.load_txt(file_path)
                else:
                    chunks = self.load_with_unstructured(file_path)

                all_text_chunks.extend(chunks)

        if not all_text_chunks:
            logger.warning("No documents found or loaded")
            return []

        # Convert to Document objects and split
        documents = []
        for chunk in all_text_chunks:
            doc = Document(
                page_content=chunk['text'],
                metadata={
                    'source': chunk['source'],
                    'page': chunk['page'],
                    'type': chunk['type']
                }
            )
            documents.append(doc)

        # Split documents into chunks
        logger.info(f"Splitting {len(documents)} documents into chunks")
        split_docs = self.text_splitter.split_documents(documents)

        logger.info(f"Created {len(split_docs)} document chunks")
        return split_docs

    def create_vector_store(self, documents: List[Document]) -> None:
        """
        Create and save vector store from documents.

        Args:
            documents: List of Document objects
        """
        if not documents:
            logger.error("No documents provided for vector store creation")
            return

        logger.info(f"Creating vector store with {len(documents)} documents")

        if self.use_faiss:
            # Create FAISS vector store
            self.vector_store = FAISS.from_documents(documents, self.embeddings)

            # Save vector store
            vector_store_path = self.embeddings_dir / "faiss_index"
            self.vector_store.save_local(str(vector_store_path))
            logger.info(f"FAISS vector store saved to {vector_store_path}")
        else:
            # Alternative: ChromaDB
            import chromadb
            from langchain.vectorstores import Chroma

            self.vector_store = Chroma.from_documents(
                documents,
                self.embeddings,
                persist_directory=str(self.embeddings_dir / "chroma_db")
            )
            logger.info(f"ChromaDB vector store saved to {self.embeddings_dir / 'chroma_db'}")

        # Save documents metadata
        metadata_path = self.embeddings_dir / "documents_metadata.pkl"
        with open(metadata_path, 'wb') as f:
            pickle.dump(documents, f)
        logger.info(f"Documents metadata saved to {metadata_path}")

        self.documents = documents

    def load_vector_store(self) -> bool:
        """
        Load existing vector store.

        Returns:
            True if vector store was loaded successfully, False otherwise
        """
        try:
            if self.use_faiss:
                vector_store_path = self.embeddings_dir / "faiss_index"
                if vector_store_path.exists():
                    self.vector_store = FAISS.load_local(
                        str(vector_store_path),
                        self.embeddings,
                        allow_dangerous_deserialization=True
                    )
                    logger.info(f"FAISS vector store loaded from {vector_store_path}")
                else:
                    logger.warning("FAISS vector store not found")
                    return False
            else:
                chroma_path = self.embeddings_dir / "chroma_db"
                if chroma_path.exists():
                    from langchain.vectorstores import Chroma
                    self.vector_store = Chroma(
                        persist_directory=str(chroma_path),
                        embedding_function=self.embeddings
                    )
                    logger.info(f"ChromaDB vector store loaded from {chroma_path}")
                else:
                    logger.warning("ChromaDB vector store not found")
                    return False

            # Load documents metadata
            metadata_path = self.embeddings_dir / "documents_metadata.pkl"
            if metadata_path.exists():
                with open(metadata_path, 'rb') as f:
                    self.documents = pickle.load(f)
                logger.info(f"Loaded {len(self.documents)} documents from metadata")

            return True

        except Exception as e:
            logger.error(f"Error loading vector store: {str(e)}")
            return False

    def retrieve(self, query: str, top_k: int = 3) -> List[Dict[str, Any]]:
        """
        Retrieve relevant documents for a given query.

        Args:
            query: Search query
            top_k: Number of documents to retrieve

        Returns:
            List of retrieved documents with relevance scores
        """
        if not self.vector_store:
            logger.error("Vector store not initialized")
            return []

        try:
            # Search for similar documents
            docs_with_scores = self.vector_store.similarity_search_with_score(
                query,
                k=top_k
            )

            results = []
            for doc, score in docs_with_scores:
                result = {
                    'content': doc.page_content,
                    'metadata': doc.metadata,
                    'score': float(score),
                    'query': query,
                    'timestamp': datetime.now().isoformat()
                }
                results.append(result)

            logger.info(f"Retrieved {len(results)} documents for query: {query[:50]}...")
            return results

        except Exception as e:
            logger.error(f"Error during retrieval: {str(e)}")
            return []

    def initialize_or_load(self, force_rebuild: bool = False) -> None:
        """
        Initialize vector store from documents or load existing one.

        Args:
            force_rebuild: Whether to force rebuild the vector store
        """
        if force_rebuild or not self.load_vector_store():
            logger.info("Building new vector store")
            documents = self.load_documents()
            if documents:
                self.create_vector_store(documents)
            else:
                logger.error("No documents found to build vector store")
        else:
            logger.info("Using existing vector store")

    def get_document_stats(self) -> Dict[str, Any]:
        """
        Get statistics about loaded documents.

        Returns:
            Dictionary with document statistics
        """
        if not self.documents:
            return {"error": "No documents loaded"}

        # Calculate statistics
        sources = [doc.metadata.get('source', 'unknown') for doc in self.documents]
        doc_types = [doc.metadata.get('type', 'unknown') for doc in self.documents]

        stats = {
            'total_documents': len(self.documents),
            'total_chars': sum(len(doc.page_content) for doc in self.documents),
            'unique_sources': len(set(sources)),
            'source_distribution': pd.Series(sources).value_counts().to_dict(),
            'type_distribution': pd.Series(doc_types).value_counts().to_dict(),
            'avg_chunk_length': np.mean([len(doc.page_content) for doc in self.documents])
        }

        return stats


if __name__ == "__main__":
    # Example usage
    retriever = GISDocumentRetriever()
    retriever.initialize_or_load()

    # Test retrieval
    query = "What is NDVI and how is it calculated?"
    results = retriever.retrieve(query, top_k=3)

    print(f"\nQuery: {query}")
    print(f"Retrieved {len(results)} documents:")
    for i, result in enumerate(results):
        print(f"\n--- Result {i+1} ---")
        print(f"Score: {result['score']:.4f}")
        print(f"Source: {result['metadata']['source']}")
        print(f"Content: {result['content'][:200]}...")